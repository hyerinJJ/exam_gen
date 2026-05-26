import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 색상 시스템 ───────────────────────────────────────────────
BLACK    = "000000"
DARK     = "1A1A1A"
MID      = "555555"
LIGHT    = "AAAAAA"
RULE     = "CCCCCC"
OFFWHITE = "F7F7F7"
WHITE    = "FFFFFF"

_PAGE_BREAK_KEYWORDS = ("1개", "한 페이지", "페이지당 1", "하나당", "1문제")
_CONTENT_W = 16.0  # A4(21cm) - 여백 2.5cm×2

TYPE_KO = {
    "short": "단답형", "essay": "에세이형", "application": "응용형",
    "list": "나열형", "order": "순서형",
}


# ── 기본 유틸 ─────────────────────────────────────────────────

def _rgb(h: str) -> RGBColor:
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _dxa(cm: float) -> str:
    return str(round(cm * 566.929))


def _set_font(run, size: int = 12, bold: bool = False,
              color: str = None, italic: bool = False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = _rgb(color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _set_margins(doc: Document, cm: float = 2.5):
    s = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, attr, Cm(cm))


def _page_break(doc):
    p = doc.add_paragraph()
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p._p.append(r)


# ── XML 직접 조작 유틸 ────────────────────────────────────────

def set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_border(cell, sides: dict):
    """sides 예: {"left": {"val":"single","sz":"20","color":"000000"}}"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in sides.items():
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   attrs.get("val",   "single"))
        b.set(qn("w:sz"),    attrs.get("sz",    "4"))
        b.set(qn("w:space"), attrs.get("space", "0"))
        b.set(qn("w:color"), attrs.get("color", BLACK))
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_horizontal_line(doc, color: str = BLACK, sz: str = "12"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── 테이블 공통 유틸 ──────────────────────────────────────────

def _make_table(doc, rows, cols, width_cm=None, center=True):
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    if center:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        tblPr.append(jc)
    if width_cm is not None:
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), _dxa(width_cm))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    return table


def _set_row_height(row, cm: float):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), _dxa(cm))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _set_table_borders(table, val="single", sz="4", color=BLACK):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_cell_width(cell, cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), _dxa(cm))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_margin(cell, top=80, bottom=80, left=120, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _cp(cell, text="", size=10, bold=False, color=DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=False,
        sb=0, sa=0, indent_cm=0, first=True):
    """셀 내 문단 추가. first=True 이면 기존 첫 문단 사용."""
    p = cell.paragraphs[0] if (first and cell.paragraphs) else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


# ── 표지 구성 요소 ────────────────────────────────────────────

def _logo_table(doc, school: str):
    table = _make_table(doc, 1, 1, width_cm=4.0)
    _set_row_height(table.rows[0], 3.0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, OFFWHITE)
    set_cell_border(cell, {
        "top":    {"val": "single", "sz": "4", "color": RULE},
        "left":   {"val": "single", "sz": "4", "color": RULE},
        "bottom": {"val": "single", "sz": "4", "color": RULE},
        "right":  {"val": "single", "sz": "4", "color": RULE},
    })
    _cp(cell, "LOGO", size=13, bold=True, color=LIGHT, sb=24, sa=4)
    _cp(cell, school, size=9, color=LIGHT, first=False, sa=4)


def _badge_table(doc, text: str):
    table = _make_table(doc, 1, 1, width_cm=5.0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    _cp(cell, text, size=13, bold=True, color=WHITE, sb=6, sa=6)


def _info_table(doc, fmt: dict):
    table = _make_table(doc, 2, 3, width_cm=_CONTENT_W, center=False)
    _set_table_borders(table, sz="4", color=RULE)
    col_w = _CONTENT_W / 3
    headers = ["담당교수", "시험일시 및 제한시간", "학기"]
    exam_dt = fmt.get("exam_date") or ""
    tl = fmt.get("time_limit") or ""
    dt_str = f"{exam_dt} / {tl}".strip(" /") if (exam_dt or tl) else ""
    data = [fmt.get("professor") or "", dt_str, fmt.get("semester") or ""]
    for i, (hdr, val) in enumerate(zip(headers, data)):
        hc = table.cell(0, i)
        set_cell_shading(hc, DARK)
        _set_cell_width(hc, col_w)
        _cp(hc, hdr, size=9, bold=True, color=WHITE, sb=4, sa=4)
        dc = table.cell(1, i)
        set_cell_shading(dc, OFFWHITE)
        _set_cell_width(dc, col_w)
        _cp(dc, val, size=9, color=DARK, sb=4, sa=4)


def _examinee_table(doc):
    table = _make_table(doc, 2, 4, width_cm=_CONTENT_W, center=False)
    _set_table_borders(table, sz="4", color=RULE)
    col_widths = [3.0, 5.0, 3.0, 5.0]
    rows_data = [
        [("학번", True), ("", False), ("이름", True), ("", False)],
        [("학과", True), ("", False), ("총점", True), ("         / 100", False)],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, (text, is_header) in enumerate(row_data):
            cell = table.cell(ri, ci)
            _set_cell_width(cell, col_widths[ci])
            if is_header:
                set_cell_shading(cell, DARK)
                _cp(cell, text, size=9, bold=True, color=WHITE, sb=6, sa=6)
            else:
                set_cell_shading(cell, WHITE)
                _cp(cell, text, size=9, color=DARK, sb=6, sa=6)


def _oath_table(doc):
    table = _make_table(doc, 1, 1, width_cm=_CONTENT_W, center=False)
    _set_table_borders(table, sz="8", color=DARK)
    cell = table.cell(0, 0)
    set_cell_shading(cell, OFFWHITE)
    _set_cell_margin(cell, top=120, bottom=120, left=180, right=180)
    _cp(cell, "— 정직 서약 —", size=11, bold=True, color=DARK, sb=4, sa=4)
    _cp(cell, "본인은 이 시험에서 어떠한 부정행위도 하지 않을 것을 서약합니다.",
        size=9, color=DARK, first=False, sb=2, sa=2)
    _cp(cell,
        "I pledge that I will not engage in any form of academic dishonesty during this exam.",
        size=9, italic=True, color=MID, first=False, sb=2, sa=4)
    _cp(cell, "서명 (Signature) : ___________________",
        size=9, color=DARK, first=False, sb=4, sa=4)


def _add_cover_page(doc: Document, fmt: dict):
    school      = fmt.get("school")      or "○○대학교"
    department  = fmt.get("department")  or ""
    title       = fmt.get("title")       or "시험지"
    english_name = fmt.get("english_name") or ""
    course_info = fmt.get("course_info") or "시험"

    doc.add_paragraph()
    _logo_table(doc, school)
    doc.add_paragraph()

    add_horizontal_line(doc, color=BLACK, sz="18")
    doc.add_paragraph()

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s.paragraph_format.space_after = Pt(2)
    _set_font(p_s.add_run(school), size=11, color=MID)

    if department:
        p_d = doc.add_paragraph()
        p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_d.paragraph_format.space_after = Pt(4)
        _set_font(p_d.add_run(department), size=9, color=LIGHT)

    add_horizontal_line(doc, color=RULE, sz="4")
    doc.add_paragraph()

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_after = Pt(6)
    _set_font(p_t.add_run(title), size=22, bold=True, color=BLACK)

    if english_name:
        p_e = doc.add_paragraph()
        p_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_e.paragraph_format.space_after = Pt(8)
        _set_font(p_e.add_run(english_name), size=11, italic=True, color=MID)

    doc.add_paragraph()
    _badge_table(doc, course_info)
    doc.add_paragraph()

    add_horizontal_line(doc, color=RULE, sz="4")
    doc.add_paragraph()

    _info_table(doc, fmt)
    doc.add_paragraph()

    _examinee_table(doc)
    doc.add_paragraph()

    _oath_table(doc)
    doc.add_paragraph()

    add_horizontal_line(doc, color=BLACK, sz="8")

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_before = Pt(4)
    _set_font(p_note.add_run(
        "시험지를 넘기기 전, 위 정보를 모두 기재하였는지 확인하십시오."),
        size=9, italic=True, color=LIGHT)

    _page_break(doc)


# ── 헤더 / 푸터 ───────────────────────────────────────────────

def _fld_run(fld_type: str):
    r = OxmlElement("w:r")
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), fld_type)
    r.append(fc)
    return r


def _instr_run(text: str):
    r = OxmlElement("w:r")
    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = text
    r.append(el)
    return r


def _add_exam_header(doc: Document, left_text: str, right_text: str,
                     different_first: bool = False):
    section = doc.sections[0]
    section.different_first_page_header_footer = different_first
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(4)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), BLACK)
    pBdr.append(bot)
    pPr.append(pBdr)

    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), _dxa(_CONTENT_W))
    tabs.append(tab)
    pPr.append(tabs)

    _set_font(p.add_run(left_text), size=9, color=DARK)
    _set_font(p.add_run("\t"), size=9)
    _set_font(p.add_run(right_text), size=9, color=DARK)


def _add_exam_footer(doc: Document, extra: str = "무단 복제 금지"):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), RULE)
    pBdr.append(top)
    pPr.append(pBdr)

    _set_font(p.add_run("페이지 "), size=9, color=LIGHT)
    for node in (_fld_run("begin"), _instr_run(" PAGE "),
                 _fld_run("separate"), _fld_run("end")):
        p._p.append(node)
    _set_font(p.add_run(" / "), size=9, color=LIGHT)
    for node in (_fld_run("begin"), _instr_run(" NUMPAGES "),
                 _fld_run("separate"), _fld_run("end")):
        p._p.append(node)
    _set_font(p.add_run(f"  |  {extra}"), size=9, color=LIGHT)


# ── 시험지 본문 구성 요소 ─────────────────────────────────────

def _add_section_header(doc, section_name: str, score_note: str):
    doc.add_paragraph()
    table = _make_table(doc, 1, 1, width_cm=_CONTENT_W, center=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, DARK)
    set_cell_border(cell, {
        "left":   {"val": "single", "sz": "24", "color": BLACK},
        "top":    {"val": "nil",    "sz": "0",  "color": "auto"},
        "right":  {"val": "nil",    "sz": "0",  "color": "auto"},
        "bottom": {"val": "nil",    "sz": "0",  "color": "auto"},
    })
    _set_cell_margin(cell, top=60, bottom=60, left=120, right=80)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(p.add_run(f"  ▶  {section_name}  |  {score_note}"),
              size=10, bold=True, color=WHITE)
    doc.add_paragraph()


def _add_answer_box(doc, question: dict):
    q_type = question.get("type", "short")
    if q_type == "short":
        row_count = 1
    elif q_type == "application":
        row_count = 10
    elif q_type in ("list", "order"):
        row_count = question.get("count") or _detect_list_count(question.get("question", ""))
    else:
        row_count = 8
    table = _make_table(doc, row_count, 1, width_cm=_CONTENT_W, center=False)
    _set_table_borders(table, val="single", sz="4", color=RULE)
    for row in table.rows:
        _set_row_height(row, 1.0)
        set_cell_shading(row.cells[0], OFFWHITE)
    doc.add_paragraph()


# ── 기존 로직 유지 ────────────────────────────────────────────

def _interpret_format(plan: dict) -> dict:
    if not plan:
        return {
            "title": None, "course_info": None, "professor": None,
            "page_break_per_question": False, "school": None,
            "department": None, "english_name": None,
            "exam_date": None, "time_limit": None, "semester": None,
        }
    layout = str(plan.get("레이아웃") or "").lower()
    page_break = any(kw in layout for kw in _PAGE_BREAK_KEYWORDS)
    return {
        "title":       plan.get("시험제목") or plan.get("과목명") or None,
        "course_info": plan.get("시험종류") or None,
        "professor":   plan.get("담당교수") or None,
        "page_break_per_question": page_break,
        "school":      plan.get("학교") or None,
        "department":  plan.get("학과") or None,
        "english_name": plan.get("영문명") or None,
        "exam_date":   plan.get("시험일시") or None,
        "time_limit":  plan.get("제한시간") or None,
        "semester":    plan.get("학기") or None,
    }


def _detect_list_count(question_text: str) -> int:
    m = re.search(r'(\d+)\s*(?:가지|개|단계|항목|종류)', question_text)
    if m:
        return min(int(m.group(1)), 8)
    return 3


# ── 메인 함수 ─────────────────────────────────────────────────

def save_exam_docx(questions: list, output_path: str, plan: dict = None) -> None:
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    fmt = _interpret_format(plan or {})

    doc = Document()
    _set_margins(doc)

    dept  = fmt.get("department") or ""
    title = fmt.get("title") or "시험지"
    prof  = fmt.get("professor") or ""
    left_hdr  = f"{dept} | {title}" if dept else title
    right_hdr = f"담당교수: {prof}" if prof else ""
    _add_exam_header(doc, left_hdr, right_hdr, different_first=True)
    _add_exam_footer(doc, "무단 복제 금지")
    _add_cover_page(doc, fmt)

    n = len(questions)
    points_per_q = round(100 / n) if n > 0 else 10
    page_break = fmt["page_break_per_question"]
    prev_type = None

    for i, q in enumerate(questions, start=1):
        q_type    = q.get("type", "short")
        q_type_ko = TYPE_KO.get(q_type, q_type)

        if q_type != prev_type:
            _add_section_header(doc, q_type_ko, f"각 문항 {points_per_q}점")
            prev_type = q_type

        p_num = doc.add_paragraph()
        p_num.paragraph_format.space_before = Pt(10)
        p_num.paragraph_format.space_after = Pt(6)
        _set_font(p_num.add_run(f"Q{i}."), size=11, bold=True, color=BLACK)
        _set_font(p_num.add_run(f"  [{points_per_q}점]"), size=10, bold=True, color=DARK)
        _set_font(p_num.add_run(f"  ({q_type_ko})"), size=9, color=LIGHT)

        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_after = Pt(8)
        _set_font(p_q.add_run(q.get("question", "")), size=11, color=DARK)

        _add_answer_box(doc, q)

        if page_break and i < n:
            _page_break(doc)

    doc.save(output_path)
    print(f"[file_writers] 시험지 저장: {output_path}")


def save_answer_key_docx(qa_pairs: list, output_path: str) -> None:
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc = Document()
    _set_margins(doc)

    _add_exam_header(doc,
                     "[학과명] | [과목명] — 모범답안 및 채점기준",
                     "배포 금지",
                     different_first=False)
    _add_exam_footer(doc, "배포 금지")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    _set_font(p_title.add_run("모범답안 및 채점기준"), size=20, bold=True, color=BLACK)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    _set_font(p_sub.add_run("[과목명] [시험종류]"), size=11, color=MID)

    add_horizontal_line(doc, color=BLACK, sz="14")
    doc.add_paragraph()

    for i, qa in enumerate(qa_pairs, start=1):
        q_type    = qa.get("type", "")
        q_type_ko = TYPE_KO.get(q_type, q_type)
        question  = qa.get("question", "")
        answer    = qa.get("answer", "")
        rubric    = qa.get("rubric", "")
        note      = qa.get("note", "")

        p_num = doc.add_paragraph()
        p_num.paragraph_format.space_before = Pt(14)
        p_num.paragraph_format.space_after = Pt(4)
        _set_font(p_num.add_run(f"Q{i}."), size=11, bold=True, color=BLACK)
        _set_font(p_num.add_run(f"  ({q_type_ko})"), size=9, color=MID)
        _set_font(p_num.add_run(f"  {question}"), size=10, color=DARK)

        # ── 1. 모범답안 박스 (왼쪽 테두리만) ──────────────────
        tbl_a = _make_table(doc, 1, 1, width_cm=_CONTENT_W, center=False)
        c_a = tbl_a.cell(0, 0)
        set_cell_shading(c_a, OFFWHITE)
        set_cell_border(c_a, {
            "left":   {"val": "single", "sz": "20", "color": DARK},
            "top":    {"val": "nil",    "sz": "0",  "color": "auto"},
            "right":  {"val": "nil",    "sz": "0",  "color": "auto"},
            "bottom": {"val": "nil",    "sz": "0",  "color": "auto"},
        })
        _set_cell_margin(c_a, top=80, bottom=80, left=180, right=80)
        _cp(c_a, "▷  모범답안", size=9, bold=True, color=DARK,
            align=WD_ALIGN_PARAGRAPH.LEFT, sb=2, sa=4)
        p_ans = c_a.add_paragraph()
        p_ans.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_ans.paragraph_format.left_indent = Cm(0.5)
        p_ans.paragraph_format.space_after = Pt(2)
        _set_font(p_ans.add_run(answer), size=10, color=DARK)

        doc.add_paragraph()

        # ── 2. 채점기준 테이블 (배점 : 채점기준) ──
        rub_w = [3.0, 13.0]
        tbl_r = _make_table(doc, 2, 2, width_cm=_CONTENT_W, center=False)
        _set_table_borders(tbl_r, sz="4", color=RULE)
        for ri, row in enumerate(tbl_r.rows):
            for ci, cell in enumerate(row.cells):
                _set_cell_width(cell, rub_w[ci])
                if ri == 0:
                    set_cell_shading(cell, DARK)
                    hdr_texts = ["배점", "채점기준"]
                    _cp(cell, hdr_texts[ci], size=9, bold=True, color=WHITE,
                        sb=3, sa=3)
                else:
                    set_cell_shading(cell, OFFWHITE)
                    data = [qa.get("score", ""), rubric]
                    align = (WD_ALIGN_PARAGRAPH.CENTER if ci == 0
                             else WD_ALIGN_PARAGRAPH.LEFT)
                    _cp(cell, str(data[ci]), size=9, color=DARK,
                        align=align, sb=4, sa=4)

        doc.add_paragraph()

        # ── 3. 유의사항 박스 (있을 때만) ──────────────────────
        if note:
            tbl_n = _make_table(doc, 1, 1, width_cm=_CONTENT_W, center=False)
            _set_table_borders(tbl_n, sz="4", color=RULE)
            c_n = tbl_n.cell(0, 0)
            set_cell_shading(c_n, OFFWHITE)
            _set_cell_margin(c_n, top=80, bottom=80, left=120, right=80)
            _cp(c_n, "※ 채점 시 유의사항", size=9, bold=True, color=DARK,
                align=WD_ALIGN_PARAGRAPH.LEFT, sb=2, sa=2)
            p_n = c_n.add_paragraph()
            p_n.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_n.paragraph_format.space_after = Pt(2)
            _set_font(p_n.add_run(note), size=9, color=DARK)
            doc.add_paragraph()

    doc.save(output_path)
    print(f"[file_writers] 답안지 저장: {output_path}")
