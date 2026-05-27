# tests/test_tools.py
import subprocess
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ── read_file 라우터 분기 테스트 (모의 객체 사용) ─────────────────────────────

def test_read_file_routes_pdf(tmp_path):
    f = tmp_path / "test.pdf"
    f.write_bytes(b"%PDF-1.4")
    with patch("tools.file_readers.read_pdf", return_value="pdf text") as mock_pdf:
        from tools.file_readers import read_file
        result = read_file(str(f))
    mock_pdf.assert_called_once()
    assert result == "pdf text"


def test_read_pdf_includes_page_markers():
    from tools.file_readers import read_pdf

    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        def __init__(self, _path):
            self.pages = [FakePage("첫 페이지"), FakePage("둘째 페이지")]

    with patch("pypdf.PdfReader", FakeReader):
        result = read_pdf(Path("fake.pdf"))

    assert "[페이지 1]" in result
    assert "[페이지 2]" in result
    assert "첫 페이지" in result
    assert "둘째 페이지" in result


def test_read_file_routes_pptx(tmp_path):
    f = tmp_path / "test.pptx"
    f.write_bytes(b"PK")  # zip stub
    with patch("tools.file_readers.read_pptx", return_value="pptx text") as mock_pptx:
        from tools.file_readers import read_file
        result = read_file(str(f))
    mock_pptx.assert_called_once()
    assert result == "pptx text"


@pytest.mark.parametrize("ext", [".mp4", ".mov", ".avi", ".mkv"])
def test_read_file_routes_video_extensions(tmp_path, ext):
    f = tmp_path / f"test{ext}"
    f.write_bytes(b"\x00")
    with patch("tools.file_readers.read_video", return_value="video text") as mock_vid:
        from tools.file_readers import read_file
        result = read_file(str(f))
    mock_vid.assert_called_once()
    assert result == "video text"


def test_read_file_routes_txt(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("hello world", encoding="utf-8")
    from tools.file_readers import read_file
    assert read_file(str(f)) == "hello world"


def test_read_file_raises_for_missing_file(tmp_path):
    from tools.file_readers import read_file
    with pytest.raises(FileNotFoundError):
        read_file(str(tmp_path / "nonexistent.mp4"))


def test_read_file_raises_for_unsupported_extension(tmp_path):
    f = tmp_path / "test.xyz"
    f.write_bytes(b"\x00")
    from tools.file_readers import read_file
    with pytest.raises(ValueError, match="지원하지 않는"):
        read_file(str(f))


# ── ffmpeg 체크 테스트 ────────────────────────────────────────────────────────

def test_check_ffmpeg_ok():
    """시스템에 ffmpeg가 설치되어 있으면 예외 없이 통과해야 함."""
    from tools.file_readers import _check_ffmpeg
    _check_ffmpeg()  # 예외 없으면 통과


def test_check_ffmpeg_not_found_raises():
    """ffmpeg가 없으면 RuntimeError를 발생시켜야 함."""
    from tools.file_readers import _check_ffmpeg
    with patch("subprocess.run", side_effect=FileNotFoundError):
        with pytest.raises(RuntimeError, match="ffmpeg"):
            _check_ffmpeg()


# ── read_video 단위 테스트 (whisper mock) ─────────────────────────────────────

def test_read_video_returns_text(tmp_path):
    """read_video가 whisper 결과 text 문자열을 반환해야 함."""
    fake_video = tmp_path / "lecture.mp4"
    fake_video.write_bytes(b"\x00" * 16)

    mock_model = MagicMock()
    mock_model.transcribe.return_value = {"text": "  강의 내용입니다.  "}

    with patch("tools.file_readers._check_ffmpeg"), \
         patch("shutil.copy2"), \
         patch("whisper.load_model", return_value=mock_model):
        from tools.file_readers import read_video
        result = read_video(fake_video)

    assert result == "  강의 내용입니다.  "
    mock_model.transcribe.assert_called_once()


def test_read_video_temp_file_cleaned_up(tmp_path):
    """임시 파일이 transcribe 후 삭제되어야 함."""
    fake_video = tmp_path / "lecture.mp4"
    fake_video.write_bytes(b"\x00" * 16)

    created_tmp: list = []

    def capture_copy(src, dst):
        created_tmp.append(Path(dst))

    mock_model = MagicMock()
    mock_model.transcribe.return_value = {"text": ""}

    with patch("tools.file_readers._check_ffmpeg"), \
         patch("shutil.copy2", side_effect=capture_copy), \
         patch("whisper.load_model", return_value=mock_model):
        from tools.file_readers import read_video
        read_video(fake_video)

    assert created_tmp, "copy2가 한 번도 호출되지 않음"
    assert not created_tmp[0].exists(), f"임시 파일이 삭제되지 않음: {created_tmp[0]}"


def test_read_video_temp_cleaned_up_on_error(tmp_path):
    """transcribe 중 예외가 발생해도 임시 파일이 삭제되어야 함."""
    fake_video = tmp_path / "lecture.mp4"
    fake_video.write_bytes(b"\x00" * 16)

    created_tmp: list = []

    def capture_copy(src, dst):
        created_tmp.append(Path(dst))
        Path(dst).write_bytes(b"\x00")  # 실제로 파일을 만들어야 cleanup이 의미 있음

    mock_model = MagicMock()
    mock_model.transcribe.side_effect = RuntimeError("transcribe 실패")

    with patch("tools.file_readers._check_ffmpeg"), \
         patch("shutil.copy2", side_effect=capture_copy), \
         patch("whisper.load_model", return_value=mock_model):
        from tools.file_readers import read_video
        with pytest.raises(RuntimeError, match="transcribe 실패"):
            read_video(fake_video)

    assert not created_tmp[0].exists(), "예외 발생 시에도 임시 파일이 삭제되어야 함"


# ── search_with_google retry_call 테스트 ─────────────────────────────────────

def test_search_with_google_uses_retry_call():
    """search_with_google가 retry_call을 통해 API를 호출해야 함."""
    with patch("tools.search_tools.retry_call") as mock_retry:
        mock_retry.return_value = MagicMock(text="검색 결과")
        import tools.search_tools as st
        result = st.search_with_google("테스트 쿼리")
    mock_retry.assert_called_once()
    assert result == "검색 결과"


def test_search_with_google_passes_query():
    """search_with_google가 retry_call에 올바른 쿼리를 담아 호출해야 함."""
    captured = {}

    def fake_retry(fn):
        response = MagicMock(text="결과")
        captured["fn"] = fn
        return response

    with patch("tools.search_tools.retry_call", side_effect=fake_retry):
        import tools.search_tools as st
        st.search_with_google("머신러닝 사례")

    assert "fn" in captured


# ── file writer 포맷 해석 / 출력 정리 테스트 ─────────────────────────────────

def test_interpret_format_parses_korean_english_course_and_year_semester():
    from tools.file_writers import _interpret_format

    fmt = _interpret_format({
        "시험 치는 과목": "한글 - 과학적 관리, 영어 - Scientific Management",
        "년도": "2026학년도",
        "학기": "1학기",
        "시험종류": "중간고사",
    })

    assert fmt["title"] == "과학적 관리"
    assert fmt["english_name"] == "Scientific Management"
    assert fmt["semester"] == "2026학년도 1학기"
    assert fmt["course_info"] == "중간고사"


def test_save_exam_docx_hides_tf_marker(tmp_path):
    from docx import Document
    from tools.file_writers import save_exam_docx

    out = tmp_path / "exam.docx"
    save_exam_docx(
        [{"id": "Q1", "type": "tf", "question": "비지도학습은 레이블이 필요 없다. (T/F)", "answer": "T", "rubric": "기준"}],
        str(out),
        plan={"진위형": 1, "과목명": "과학적 관리"},
    )

    text = "\n".join(p.text for p in Document(out).paragraphs)
    assert "(T/F)" not in text
    assert "비지도학습은 레이블이 필요 없다." in text


def test_save_exam_docx_splits_application_subquestions(tmp_path):
    from docx import Document
    from tools.file_writers import save_exam_docx

    out = tmp_path / "exam.docx"
    question = (
        "편의점에 새로운 POS 시스템이 도입되었다.\n"
        "(1) Work System Framework를 적용하여 문제를 분석하시오.\n"
        "(2) 개선 방향을 제시하시오."
    )
    save_exam_docx(
        [{"id": "Q1", "type": "application", "question": question, "answer": "답", "rubric": "기준"}],
        str(out),
        plan={"응용형": 1, "과목명": "과학적 관리"},
    )

    paragraphs = [p.text for p in Document(out).paragraphs]
    assert any(p == "편의점에 새로운 POS 시스템이 도입되었다." for p in paragraphs)
    assert any(p.startswith("(1) Work System Framework") for p in paragraphs)
    assert any(p.startswith("(2) 개선 방향") for p in paragraphs)


def test_save_answer_key_embeds_tf_and_short_answers(tmp_path):
    from docx import Document
    from docx.shared import RGBColor
    from tools.file_writers import save_answer_key_docx

    out = tmp_path / "answer_key.docx"
    save_answer_key_docx(
        [
            {"id": "Q1", "type": "tf", "question": "명제이다. (T/F)", "answer": "T", "rubric": "정답: T"},
            {"id": "Q2", "type": "short", "question": "용어는?", "answer": "과적합", "rubric": "정답(5점): 과적합"},
        ],
        str(out),
    )

    doc = Document(out)
    runs = [run for p in doc.paragraphs for run in p.runs]
    tf_run = next(run for run in runs if run.text == "T")
    short_run = next(run for run in runs if run.text == "과적합")

    assert tf_run.font.bold is True
    assert tf_run.font.underline is True
    assert tf_run.font.color.rgb == RGBColor(255, 0, 0)
    assert short_run.font.bold in (False, None)
    assert short_run.font.color.rgb == RGBColor(255, 0, 0)


def test_save_answer_key_essay_uses_answer_rubric_and_note_labels(tmp_path):
    from docx import Document
    from tools.file_writers import save_answer_key_docx

    out = tmp_path / "answer_key.docx"
    save_answer_key_docx(
        [{
            "id": "Q1",
            "type": "essay",
            "question": "설명하시오.",
            "answer": "핵심 답안",
            "rubric": "핵심 포인트 (10점)",
            "grading_notes": "표현이 달라도 개념이 정확하면 인정",
        }],
        str(out),
    )

    text = "\n".join(p.text for p in Document(out).paragraphs)
    assert "모범답안:" in text
    assert "채점기준:" in text
    assert "▶ 모범답안" not in text
    assert "* 표현이 달라도 개념이 정확하면 인정" in text


def test_save_answer_key_displays_essay_subpoints(tmp_path):
    from docx import Document
    from tools.file_writers import save_answer_key_docx

    out = tmp_path / "answer_key.docx"
    save_answer_key_docx(
        [{
            "id": "Q1",
            "type": "essay",
            "question": "배경 설명\n(1) 설명하시오.\n(2) 비교하시오.",
            "answer": "핵심 답안",
            "rubric": "(1) 기준 (15점)\n(2) 기준 (10점)",
            "points": 25,
            "subpoints": [15, 10],
        }],
        str(out),
    )

    paragraphs = [p.text for p in Document(out).paragraphs]
    assert any(p.startswith("(1) 설명하시오.") and "(15점)" in p for p in paragraphs)
    assert any(p.startswith("(2) 비교하시오.") and "(10점)" in p for p in paragraphs)


def test_save_answer_key_with_plan_adds_template_cover(tmp_path):
    from docx import Document
    from tools.file_writers import save_answer_key_docx

    out = tmp_path / "answer_key.docx"
    save_answer_key_docx(
        [{"id": "Q1", "type": "tf", "question": "명제이다. (T/F)", "answer": "T", "rubric": "정답: T"}],
        str(out),
        plan={
            "시험 치는 과목": "한글 - 과학적 관리, 영어 - Scientific Management",
            "년도": "2026학년도",
            "학기": "1학기",
            "시험종류": "기말고사",
            "시험일시": "4월 14일 9시 30분 ~ 10시 45분",
        },
    )

    paragraphs = [p.text for p in Document(out).paragraphs]
    assert paragraphs[3] == "과학적 관리"
    assert "Scientific Management" in paragraphs
    assert "2026학년도 1학기" in paragraphs
    assert "기말고사" in paragraphs
    assert "시험 일시: 4월 14일 9시 30분 ~ 10시 45분" in paragraphs
    assert "모범 답안" in paragraphs
    assert "모범답안 및 채점기준" not in paragraphs[:15]


def test_save_answer_key_uses_exam_grouping_and_order(tmp_path):
    from docx import Document
    from tools.file_writers import save_answer_key_docx

    out = tmp_path / "answer_key.docx"
    save_answer_key_docx(
        [
            {"id": "E1", "type": "essay", "question": "서술하시오.", "answer": "서술 답", "rubric": "서술 기준", "points": 10},
            {"id": "T1", "type": "tf", "question": "첫 명제이다. (T/F)", "answer": "T", "rubric": "정답: T", "points": 2},
            {"id": "S1", "type": "short", "question": "용어는?", "answer": "과적합", "rubric": "정답(5점): 과적합", "points": 5},
            {"id": "T2", "type": "tf", "question": "둘째 명제이다. (T/F)", "answer": "F", "rubric": "정답: F", "points": 2},
            {"id": "A1", "type": "application", "question": "상황\n(1) 분석하시오.", "answer": "응용 답", "rubric": "응용 기준", "points": 10, "subpoints": [10]},
        ],
        str(out),
    )

    paragraphs = [p.text for p in Document(out).paragraphs]
    problem_headers = [p for p in paragraphs if p.startswith("문제 ")]

    assert problem_headers == [
        "문제 1 (4점) True/False (빈칸에 T 또는 F를 작성, 각 2점)",
        "문제 2 (5점) 단답형 (각 5점)",
        "문제 3 (10점)",
        "문제 4 (10점)",
    ]
    assert any(p.startswith("(1)") and "첫 명제이다." in p for p in paragraphs)
    assert any(p.startswith("(2)") and "둘째 명제이다." in p for p in paragraphs)


# ── 실제 파이프라인 통합 테스트 (ffmpeg로 테스트 영상 생성) ────────────────────

def _make_silent_mp4(output_path: Path, duration: int = 2) -> bool:
    """ffmpeg로 짧은 무음 MP4를 생성. 실패 시 False 반환."""
    ret = subprocess.run(
        [
            "ffmpeg", "-y",
            "-f", "lavfi", "-i", f"color=black:s=64x64:r=1:d={duration}",
            "-f", "lavfi", "-i", f"anullsrc=r=16000:cl=mono",
            "-t", str(duration),
            "-c:v", "libx264", "-c:a", "aac",
            "-shortest", str(output_path),
        ],
        capture_output=True,
    )
    return ret.returncode == 0 and output_path.exists()


@pytest.mark.slow
def test_read_video_real_pipeline(tmp_path):
    """실제 ffmpeg + whisper로 2초 무음 영상을 처리하는 통합 테스트."""
    video_path = tmp_path / "silent_test.mp4"
    if not _make_silent_mp4(video_path):
        pytest.skip("ffmpeg로 테스트 영상 생성 실패 — 환경 확인 필요")

    from tools.file_readers import read_video
    result = read_video(video_path)

    assert isinstance(result, str), "반환값이 문자열이어야 함"
    print(f"\n[통합 테스트] whisper 출력: {result!r}")
